import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF extract error: {e}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # ──── Pull the tables out too ────
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"DOCX extract error: {e}")
    return text.strip()


def extract_policy_text(file_path: str, file_type: str) -> str:
    """
    Extract text according to the file type
    file_type = 'pdf' ya 'docx'
    """
    if file_type == "pdf":
        text = extract_text_from_pdf(file_path)
    elif file_type == "docx":
        text = extract_text_from_docx(file_path)
    else:
        return ""

    # ──── If the text came back empty ────
    if not text:
        print(f"Warning: No text extracted from {file_path}")
        return ""

    return text